import click
import pandas as pd
from docx import Document


def save_as_csv(data, product_name):
    """Veriyi CSV olarak kaydeder."""
    df = pd.DataFrame(data)
    filename = f"{product_name}.csv"
    df.to_csv(filename, index=False, encoding="utf-8-sig")
    click.echo(f"Veri başarıyla {filename} dosyasına kaydedildi.")


def save_as_excel(data, product_name):
    """Veriyi Excel olarak kaydeder."""
    df = pd.DataFrame(data)
    filename = f"{product_name}.xlsx"
    df.to_excel(filename, index=False)
    click.echo(f"Veri başarıyla {filename} dosyasına kaydedildi.")


def save_as_word(data, product_name):
    """Veriyi Word olarak kaydeder."""
    df = pd.DataFrame(data)
    filename = f"{product_name}.docx"
    doc = Document()
    doc.add_heading(f"{product_name.capitalize()} Fiyat Raporu", 0)

    if df.empty:
        doc.add_paragraph("Bu ürün için veri bulunamadı.")
    else:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    doc.save(filename)
    click.echo(f"Veri başarıyla {filename} dosyasına kaydedildi.")
