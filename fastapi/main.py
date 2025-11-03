import os
# Proje kök dizinini sys.path'e ekleyerek utils modüllerine erişim sağlıyoruz
import sys
from io import BytesIO, StringIO

import pandas as pd
from docx import Document

from fastapi import FastAPI, Query
from fastapi.responses import StreamingResponse

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.api_client import fetch_data
from utils.data_processor import process_data

app = FastAPI(
    title="Piyasa Fiyatları API",
    description="Belirtilen ürün için piyasa fiyatlarını çeker ve farklı formatlarda sunar.",
    version="1.0.0",
)


async def get_processed_data(product_name: str):
    api_response = await fetch_data(product_name)
    if not api_response:
        return None
    return process_data(api_response)


@app.get(
    "/search/",
    summary="Ürün Fiyatlarını Ara",
    description="Bir ürün adını kullanarak fiyatları arar ve belirtilen formatta bir dosya olarak döndürür.",
)
async def search_product(
    product_name: str = Query(
        ..., description="Aranacak ürün adı (örn: domates, zeytinyağı)"
    ),
    format: str = Query(
        "csv", enum=["csv", "excel", "word"], description="Döndürülecek dosya formatı."
    ),
):
    """
    Bu endpoint, kullanıcı tarafından sağlanan ürün adını alır, fiyat verilerini
    asenkron olarak çeker, işler ve seçilen formatta (CSV, Excel veya Word)
    bir dosya olarak geri döndürür.
    """
    data = await get_processed_data(product_name)

    if not data:
        return {"message": "Ürün için veri bulunamadı."}

    df = pd.DataFrame(data)
    filename = f"{product_name.replace(' ', '_')}_fiyatlari"

    if format == "csv":
        output = BytesIO()
        df.to_csv(output, index=False, encoding='utf-8-sig')
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}.csv"}
        )

    elif format == "excel":
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}.xlsx"},
        )

    elif format == "word":
        output = BytesIO()
        doc = Document()
        doc.add_heading(f"{product_name.capitalize()} Fiyat Raporu", 0)
        if not df.empty:
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name)
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
        else:
            doc.add_paragraph("Veri bulunamadı.")
        doc.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"},
        )


# Sunucuyu çalıştırmak için: uvicorn fastapi.main:app --reload
